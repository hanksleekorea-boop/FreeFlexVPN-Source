from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import build_app_service_plan_v2 as base


base.SOURCE = base.ROOT / "10_STATE" / "DEV_EXECUTION_PLAN_v2.0_2026-08-01.md"
base.OUTPUT = base.ROOT / "60_OUTPUTS" / "FreeFlexVPN_상세개발실행계획서_v2.0_2026-08-01.docx"
base.CORE_TITLE = "FreeFlexVPN 상세 개발 실행계획서 v2.0"
base.CORE_SUBJECT = "새 앱서비스 기획을 최단 안전 알파로 구현하는 개발 실행계획"
base.CORE_COMMENT = "구현 승인 · 알파 대기"
base.RUNNING_HEADER = "FreeFlexVPN · Alpha Development Execution Plan"


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("ALPHA DEVELOPMENT EXECUTION PLAN")
    base.set_font(run, size=10, bold=True, color=base.MINT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("FreeFlexVPN")
    base.set_font(run, size=31, bold=True, color=base.INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("상세 개발 실행계획서 v2.0")
    base.set_font(run, size=16, color=base.DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    base.add_inline(
        p,
        "새 기획의 핵심 가치를 실제 기능과 디자인 UX로 먼저 구현해 가장 빠른 안전 알파를 공개한다.",
        size=14,
        color=base.INK,
        bold=True,
    )

    rows = [
        ("기준일", "2026-08-01"),
        ("상태", "설계 완료 · 구현 승인 · 알파 대기"),
        ("TTFV", "구현 시작 후 4~6 개발시간"),
        ("A0 후보", "44~60 개발시간 · 서버 접근 후 최단 5개 작업일"),
        ("범위", "실제 무료 VPN · 순간 UX · 3잔액 · 추천 보상"),
        ("후행", "실제 금전 결제는 A1 유료 알파로 분리"),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{label}: ")
        base.set_font(run, size=9.5, bold=True, color=base.DARK_BLUE)
        run = p.add_run(value)
        base.set_font(run, size=9.5, color=base.INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    base.add_inline(p, "첫 릴리스", size=10, bold=True, color=base.MINT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    base.add_inline(
        p,
        "R3-ui-truth: 서울 허구 항목 제거 · 실제 서버 카탈로그 · 보호 상태 5단계",
        size=12,
        bold=True,
        color=base.INK,
    )

    doc.add_page_break()


base.add_cover = add_cover


if __name__ == "__main__":
    base.build()
