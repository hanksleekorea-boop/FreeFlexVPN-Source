from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "10_STATE" / "APP_SERVICE_PLAN_v2.0_2026-08-01.md"
OUTPUT = ROOT / "60_OUTPUTS" / "FreeFlexVPN_앱서비스기획서_v2.0_2026-08-01.docx"
CORE_TITLE = "FreeFlexVPN 앱서비스 기획서 v2.0"
CORE_SUBJECT = "상황형 라이트 사용자를 위한 비구독 VPN 제품 기획"
CORE_COMMENT = "제품 방향 채택 · 구현 대기"
RUNNING_HEADER = "FreeFlexVPN · Product Service Plan"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
INK = "172033"
BLUE = "176B87"
DARK_BLUE = "124559"
MINT = "19A974"
MUTED = "667085"
LIGHT = "F2F4F7"
CALLOUT = "EAF7F2"
WHITE = "FFFFFF"
RISK = "9B1C1C"


def set_font(run, size=None, bold=None, color=None, italic=None, east_asia="Malgun Gothic"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def widths_for(count):
    patterns = {
        2: [2700, 6660],
        3: [1800, 3580, 3980],
        4: [1560, 2250, 2700, 2850],
        5: [1120, 1250, 2180, 2280, 2530],
    }
    if count in patterns:
        return patterns[count]
    base = CONTENT_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_inline(paragraph, text, *, size=10.5, color=INK, bold=False):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font(run, size=size, color=color, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=size - 0.5, color=DARK_BLUE, bold=True, east_asia="Malgun Gothic")
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            set_font(run, size=size, color=BLUE)
            run.font.underline = True
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=size, color=color, bold=bold)


def new_num_id(doc, fmt, text, left=720, hanging=360):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    level.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_number(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("FreeFlexVPN v2.0  |  ")
    set_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    fld_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    r_pr.extend([color, size])
    fld_run.extend([r_pr, begin, instr, separate, text, end])
    paragraph._p.append(fld_run)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("PRODUCT SERVICE PLAN")
    set_font(run, size=10, bold=True, color=MINT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("FreeFlexVPN")
    set_font(run, size=31, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("앱서비스 기획서 v2.0")
    set_font(run, size=16, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    add_inline(p, "가끔 VPN이 필요한 순간, 구독 없이 바로 보호받고 필요한 데이터만 평생 보관해 쓰는 라이트 VPN.", size=14, color=INK, bold=True)

    rows = [
        ("기준일", "2026-08-01"),
        ("상태", "제품 방향 채택 · 구현 대기"),
        ("핵심 계약", "월 1GB 무료 · 구매/추천 데이터 무기한 · 자동결제 없음"),
        ("증거 경계", "공개 UI 있음 · 실제 VPN/결제/기기/파일럿 없음"),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{label}: ")
        set_font(run, size=9.5, bold=True, color=DARK_BLUE)
        run = p.add_run(value)
        set_font(run, size=9.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, "핵심 변화", size=10, bold=True, color=MINT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    add_inline(p, "서버 중심 VPN → 필요한 순간 해결 · 무료 안전 기본기 · 3잔액 데이터 지갑 · 양쪽 500MB 추천 파일럿", size=12, bold=True, color=INK)

    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        parts = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            rows.append(parts)
        index += 1
    return rows, index


def add_table(doc, rows):
    if not rows:
        return
    col_count = len(rows[0])
    if any(len(row) != col_count for row in rows):
        raise ValueError("inconsistent markdown table columns")
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.cell(ri, ci)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri == 0:
                shade_cell(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.08
            if ci == 0 or ri == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, size=8.6, bold=(ri == 0), color=DARK_BLUE if ri == 0 else INK)
    repeat_table_header(table.rows[0])
    set_table_geometry(table, widths_for(col_count))
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def build():
    source_text = SOURCE.read_text(encoding="utf-8")
    lines = source_text.splitlines()
    doc = Document()
    doc.core_properties.title = CORE_TITLE
    doc.core_properties.subject = CORE_SUBJECT
    doc.core_properties.author = "FreeFlexVPN"
    doc.core_properties.comments = CORE_COMMENT

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(RUNNING_HEADER)
    set_font(run, size=8.5, color=MUTED)
    set_page_number(section.footer.paragraphs[0])

    configure_styles(doc)
    bullet_id = new_num_id(doc, "bullet", "•")
    decimal_id = new_num_id(doc, "decimal", "%1.")
    add_cover(doc)

    start = next(i for i, line in enumerate(lines) if line.startswith("## 0."))
    i = start
    h1_count = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        text = raw.strip()
        if not text:
            i += 1
            continue
        if text.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if text.startswith("## "):
            h1_count += 1
            p = doc.add_paragraph(style="Heading 1")
            if h1_count > 1:
                p.paragraph_format.page_break_before = True
            add_inline(p, text[3:], size=16, bold=True, color=BLUE)
        elif text.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, text[4:], size=13, bold=True, color=BLUE)
        elif re.match(r"^- ", text):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            apply_number(p, bullet_id)
            add_inline(p, text[2:], size=10.5)
        elif re.match(r"^\d+\. ", text):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            apply_number(p, decimal_id)
            add_inline(p, re.sub(r"^\d+\. ", "", text), size=10.5)
        elif text.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, text.lstrip("> "), size=9.5, color=MUTED)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.widow_control = True
            add_inline(p, text.replace("  ", " "), size=10.5)
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"DOCX written: {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    try:
        build()
    except Exception as exc:
        print(f"BUILD FAIL: {type(exc).__name__}: {exc}", file=sys.stderr)
        raise
